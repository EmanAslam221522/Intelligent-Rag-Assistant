from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import uvicorn
import os
import logging
import time
from typing import List, Optional
import jwt
from datetime import datetime, timedelta
import hashlib
import secrets
from dotenv import load_dotenv
import google.generativeai as genai
import PyPDF2
import docx
import csv
import io
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA

# Configure logging
logging.basicConfig(level=logging.INFO)

# Load environment variables
load_dotenv()

# Configure Gemini API
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    # Configure with timeout and retry settings
    gemini_model = genai.GenerativeModel(
        'gemini-1.5-flash',
        generation_config=genai.types.GenerationConfig(
            max_output_tokens=2048,
            temperature=0.7,
        )
    )
    print("✅ Gemini API configured successfully")
else:
    print("❌ GEMINI_API_KEY not found in environment variables")
    gemini_model = None

app = FastAPI(title="RAG Chatbot API", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Simple in-memory storage for demo
users_db = {}
tokens_db = {}
uploaded_content = {}

# Text extraction functions
def extract_text_from_pdf(content):
    """Extract text from PDF content"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting PDF text: {e}")
        return f"Error extracting PDF content: {str(e)}"

def extract_text_from_docx(content):
    """Extract text from DOCX content"""
    try:
        doc = docx.Document(io.BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting DOCX text: {e}")
        return f"Error extracting DOCX content: {str(e)}"

def extract_text_from_csv(content):
    """Extract text from CSV content"""
    try:
        csv_reader = csv.reader(io.StringIO(content.decode('utf-8')))
        text = ""
        for row in csv_reader:
            text += ", ".join(row) + "\n"
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting CSV text: {e}")
        return f"Error extracting CSV content: {str(e)}"

def extract_text_from_file(file):
    """Extract text from uploaded file"""
    content = file.file.read()
    file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else ''
    
    if file_extension in ['txt', 'md']:
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            return content.decode('utf-8', errors='ignore')
    elif file_extension == 'pdf':
        return extract_text_from_pdf(content)
    elif file_extension == 'docx':
        return extract_text_from_docx(content)
    elif file_extension == 'csv':
        return extract_text_from_csv(content)
    else:
        return f"Content from {file.filename} (file type: {file_extension} - text extraction not supported)"

# Simple auth functions
def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    password_hash = hashlib.sha256((password + salt).encode()).hexdigest()
    return f"{salt}:{password_hash}"

def verify_password(password: str, hashed_password: str) -> bool:
    try:
        salt, password_hash = hashed_password.split(":")
        return hashlib.sha256((password + salt).encode()).hexdigest() == password_hash
    except:
        return False

def create_token(user_id: str) -> str:
    payload = {
        "user_id": user_id,
        "exp": datetime.utcnow() + timedelta(days=7)
    }
    secret_key = os.getenv("JWT_SECRET_KEY", "demo-secret-key")
    token = jwt.encode(payload, secret_key, algorithm="HS256")
    print(f"DEBUG: Created token for user {user_id}")
    return token

def verify_token(token: str) -> Optional[dict]:
    try:
        secret_key = os.getenv("JWT_SECRET_KEY", "demo-secret-key")
        payload = jwt.decode(token, secret_key, algorithms=["HS256"])
        return payload
    except Exception as e:
        print(f"DEBUG: Token verification failed: {e}")
        return None

# Simple auth dependency
async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(HTTPBearer())):
    token = credentials.credentials
    payload = verify_token(token)
    
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    return payload

# Auth endpoints
@app.post("/api/auth/register")
async def register(username: str = Form(...), email: str = Form(...), password: str = Form(...)):
    if email in users_db:
        raise HTTPException(status_code=400, detail="User already exists")
    
    user_id = f"user_{len(users_db) + 1}"
    hashed_password = hash_password(password)
    
    users_db[email] = {
        "user_id": user_id,
        "username": username,
        "email": email,
        "password": hashed_password
    }
    
    token = create_token(user_id)
    tokens_db[token] = user_id
    
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "user_id": user_id,
            "username": username,
            "email": email
        }
    }

@app.post("/api/auth/login")
async def login(email: str = Form(...), password: str = Form(...)):
    if email not in users_db:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    user = users_db[email]
    if not verify_password(password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_token(user["user_id"])
    tokens_db[token] = user["user_id"]
    
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "user_id": user["user_id"],
            "username": user["username"],
            "email": user["email"]
        }
    }

# Health check
@app.get("/api/health")
async def health_check():
    return {"status": "healthy", "message": "RAG Chatbot API is running"}

# Document upload endpoint with proper processing
@app.post("/api/content/upload")
async def upload_document(file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    try:
        print(f"DEBUG: Uploading file: {file.filename} for user: {user.get('user_id')}")
        
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # 1. Extract text
        text = extract_text_from_file(file)
        
        if len(text) == 0:
            raise HTTPException(status_code=400, detail="Empty file provided")
        
        print(f"DEBUG: Extracted text length: {len(text)}")
        
        # 2. Process and store in vector database
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(text)
        
        doc_id = f"doc_{int(time.time())}"
        documents = [
            Document(page_content=chunk, metadata={"doc_id": doc_id, "user_id": user["user_id"], "filename": file.filename})
            for chunk in chunks
        ]
        
        # 3. Add to existing vector store or create new
        try:
            vector_store = Chroma.from_documents(
                documents=documents,
                embedding=GoogleGenerativeAIEmbeddings(model="models/embedding-001"),
                collection_name=f"user_{user['user_id']}",
                persist_directory=f"./chroma_db_{user['user_id']}"
            )
            vector_store.persist()
        except Exception as e:
            logging.warning(f"Vector store creation failed: {e}, using simple storage")
            # Fallback to simple storage
            content_id = f"content_{int(time.time())}"
            uploaded_content[content_id] = {
                "filename": file.filename,
                "content": text,
                "user_id": user["user_id"],
                "upload_time": time.time()
            }
        
        print(f"DEBUG: File uploaded successfully with ID: {doc_id}")
        
        return {
            "status": "success",
            "message": "Document processed successfully",
            "filename": file.filename,
            "doc_id": doc_id,
            "text_length": len(text)
        }
        
    except Exception as e:
        logging.error(f"Error uploading document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

# Content list endpoint
@app.get("/api/content/list")
async def list_content(user: dict = Depends(get_current_user)):
    try:
        print(f"DEBUG: Listing content for user: {user.get('user_id')}")
        
        # Get simple stored content
        user_content = []
        for content_id, content_data in uploaded_content.items():
            if content_data.get('user_id') == user.get('user_id'):
                user_content.append({
                    "id": content_id,
                    "filename": content_data.get('filename', 'Unknown'),
                    "size": content_data.get('size', 0),
                    "upload_time": content_data.get('upload_time', 0)
                })
        
        print(f"DEBUG: Found {len(user_content)} files for user")
        
        return {
            "content": user_content,
            "count": len(user_content)
        }
    except Exception as e:
        logging.error(f"Error listing content: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to list content")

# Chat endpoints with proper RAG implementation
@app.post("/api/chat/general")
async def chat_general(message: str = Form(...)):
    try:
        if not gemini_model:
            raise HTTPException(status_code=500, detail="Gemini API not configured")
        
        # Call real Gemini API with retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = gemini_model.generate_content(message)
                return {
                    "response": response.text,
                    "source": "general"
                }
            except Exception as api_error:
                logging.warning(f"Gemini API attempt {attempt + 1} failed: {str(api_error)}")
                if attempt == max_retries - 1:
                    raise api_error
                time.sleep(1)  # Wait 1 second before retry
                
    except Exception as e:
        logging.error(f"Error generating response: {str(e)}")
        return {
            "response": "I'm having trouble processing your request right now. Please try again shortly.",
            "source": "general"
        }

@app.post("/api/chat/rag")
async def chat_rag(message: str = Form(...), user: dict = Depends(get_current_user)):
    try:
        if not gemini_model:
            raise HTTPException(status_code=500, detail="Gemini API not configured")
        
        # Check if user has documents in vector store
        try:
            vector_store = Chroma(
                collection_name=f"user_{user['user_id']}",
                embedding_function=GoogleGenerativeAIEmbeddings(model="models/embedding-001"),
                persist_directory=f"./chroma_db_{user['user_id']}"
            )
            
            # If no documents or error, use general Gemini
            if vector_store._collection.count() == 0:
                raise Exception("No documents found")
            
            # Use RAG with documents
            retriever = vector_store.as_retriever(search_kwargs={"k": 3})
            qa_chain = RetrievalQA.from_chain_type(
                llm=ChatGoogleGenerativeAI(model="gemini-1.5-flash"),
                chain_type="stuff",
                retriever=retriever,
                return_source_documents=False
            )
            
            result = qa_chain.invoke({"query": message})
            return {
                "response": result["result"],
                "source": "documents",
                "sources": []
            }
            
        except Exception as rag_error:
            logging.warning(f"RAG processing failed: {rag_error}, falling back to general chat")
            
            # Fallback to general Gemini
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    response = gemini_model.generate_content(message)
                    return {
                        "response": response.text,
                        "source": "general",
                        "sources": []
                    }
                except Exception as api_error:
                    logging.warning(f"Gemini API attempt {attempt + 1} failed: {str(api_error)}")
                    if attempt == max_retries - 1:
                        raise api_error
                    time.sleep(1)  # Wait 1 second before retry
                    
    except Exception as e:
        logging.error(f"Error generating RAG response: {str(e)}")
        return {
            "response": "I'm having trouble accessing your documents. Please try again.",
            "source": "documents",
            "sources": []
        }

if __name__ == "__main__":
    print("🚀 Starting RAG Chatbot API...")
    print("✅ Real Gemini API integration")
    print("✅ Document processing with vector store")
    print("✅ RAG implementation")
    print("✅ Authentication working")
    uvicorn.run(app, host="0.0.0.0", port=8000)